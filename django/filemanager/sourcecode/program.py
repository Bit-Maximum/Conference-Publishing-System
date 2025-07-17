import io
import os

from docx.enum.section import WD_ORIENTATION
import copy
import re

from docx import Document
import docx
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from lxml import etree

from .markup import add_blank_lines, set_align_center_paragraph, set_align_left_paragraph, fix_hanging_prepositions
from .styles import *


def compose_program(doc_out, data):
    output = docx.Document()
    set_default_style(output, doc_out)
    set_markup(output)
    for i, section in enumerate(data):
        compose_program_section(output, section)
        if i < len(data) - 1:
            output.add_page_break()
    output.save(doc_out)


def compose_program_section(output, data):
    compose_heading(output, data.get("section"), data.get("moderator", None), data.get("conference_name"))
    compose_table(output, data)
    compose_subscription(output)


def set_markup(document):
    for section in document.sections:
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width = 10692130
        section.page_height = 7560310
        section.top_margin = Cm(2.25)
        section.right_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)


def compose_heading(document, section_name, moderator: dict, conference_name: str):
    compose_section_name(document, section_name=section_name, conference_name=conference_name)
    add_blank_lines(document)

    if moderator is not None:
        compose_moderator(
                document,
                name=moderator.get("name"),
                degree=moderator.get("degree"),
                academic_title=moderator.get("academic_title"),
                job_title=moderator.get("job_title"),
                job=moderator.get("job")
            )
        add_blank_lines(document)


def compose_section_name(document, section_name, conference_name):
        p = set_align_center_paragraph(document)
        run = p.add_run(f"Научная программа секции «{section_name}»")
        header_run_style(run)
        fix_hanging_prepositions(p)

        # p = set_align_center_paragraph(document)
        # run = p.add_run(f"ХIII научно-практической конференции студентов и аспирантов")
        # header_run_style(run)

        p = set_align_center_paragraph(document)
        run = p.add_run(f"«{conference_name}»")
        header_run_style(run)


def compose_moderator(document, name: str = None, degree: str = None, academic_title: str = None, job_title: str = None, job: str = None, fix_hanging: bool = False):

    if job_title and job:
        job_string = f", {job_title} {job}"
    elif job:
        job_string = f" {job}"
    elif job_title:
        job_string = f", {job_title}"
    else:
        job_string = None

    adviser_title = ', '.join([degree, academic_title])

    tun_text = f"{adviser_title}{job_string}" if job_string else adviser_title

    p = set_align_center_paragraph(document)
    run = p.add_run(f"Модератор:\u00A0")
    moderator_run_style(run)

    run = p.add_run(f"{name}, {tun_text}")
    moderator_data_run_style(run)


def compose_table(document, data):
    table = document.add_table(0, 0)
    table.add_column(Cm(0.88))
    table.add_column(Cm(3.62))
    table.add_column(Cm(11.25))
    table.add_column(Cm(5))
    table.add_column(Cm(5))

    row = table.add_row()
    row.cells[0].paragraphs[0].text = "№"
    row.cells[1].paragraphs[0].text = "ФИО докладчика(ов)"
    row.cells[2].paragraphs[0].text = "Тема доклада"
    row.cells[3].paragraphs[0].text = "ФИО научного руководителя"
    row.cells[4].paragraphs[0].text = "Уч. степень, уч. звание, должность научного руководителя"

    for i, item in enumerate(data.get("articles")):
        row = table.add_row()
        row.cells[0].paragraphs[0].text = f"{i + 1}."
        row.cells[1].paragraphs[0].text = item.get("authors")
        row.cells[2].paragraphs[0].text = item.get("title")
        row.cells[3].paragraphs[0].text = item.get("adviser").get("name")
        row.cells[4].paragraphs[0].text = item.get("adviser").get("info")

    set_style_table(table)
    add_blank_lines(document, 2)


def compose_subscription(document):
    p = set_align_left_paragraph(document)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(20))
    run = p.add_run("Должность_________________\t__________________ФИО")
    moderator_run_style(run)
    add_blank_lines(document)

    p = set_align_left_paragraph(document)
    run = p.add_run("Дата_______________________")
    moderator_run_style(run)


def header_run_style(run):
    run.font.bold = True
    run.font.color.rgb = None
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    run.font.all_caps = False


def moderator_run_style(run):
    run.font.bold = False
    run.font.color.rgb = None
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    run.font.all_caps = False


def moderator_data_run_style(run):
    moderator_run_style(run)
    run.italic = True


def set_style_table(table):
    table.style = 'Table Grid'
    table.autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = table.rows[0]
    for cell in headers.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph.paragraph_format.line_spacing = 1
            for run in paragraph.runs:
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
                run.font.bold = True

    for row in table.rows[1:]:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                paragraph.paragraph_format.line_spacing = 1
                for run in paragraph.runs:
                    run.font.size = Pt(12)
                    run.font.name = 'Times New Roman'
