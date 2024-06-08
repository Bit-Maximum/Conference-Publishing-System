from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from copy import deepcopy

from .styles import reset_paragraph_style


def set_markup(document):
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.right_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)


def set_collection_markup(document):
    section = document.sections[0]
    section.top_margin = Cm(1.7)
    section.right_margin = Cm(2)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2)


def set_default_paragraph(document, line_spacing=1.15):
    p = document.add_paragraph()
    reset_paragraph_style(p)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.5)
    return p


def set_source_paragraph(document):
    p = document.add_paragraph()
    reset_paragraph_style(p)
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    return p


def set_align_right_paragraph(document):
    p = document.add_paragraph()
    reset_paragraph_style(p)
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    return p


def set_title_paragraph(document):
    p = document.add_paragraph()
    reset_paragraph_style(p)
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return p


def set_wide_title_paragraph(document):
    p = set_title_paragraph(document)
    p.paragraph_format.line_spacing = 1.30
    return p


def set_align_center_paragraph(document, line_spacing=1.15):
    p = document.add_paragraph()
    reset_paragraph_style(p)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return p


def set_align_left_paragraph(document, line_spacing=1.15):
    p = document.add_paragraph()
    reset_paragraph_style(p)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return p


def set_underline_paragraph(document):
    p = document.add_paragraph()
    reset_paragraph_style(p)
    p.paragraph_format.line_spacing = Pt(6)


def insert_footnote(paragraph, text):
    paragraph.add_footnote(text)
    paragraph.runs[-1].font.size = Pt(12)
    paragraph.runs[-1].font.name = "Times New Roman"
    paragraph.runs[-1].font.superscript = True
    paragraph.runs[-1].font.color.rgb = None
    paragraph.runs[-1].font.bold = None
    paragraph.runs[-1].italic = None
    paragraph.runs[-1].font.all_caps = False


def is_paragraph_empty(paragraph):
    text = paragraph.text.strip()
    return not text


def get_table_copy(template):
    tbl = template._tbl
    return deepcopy(tbl)


def direct_insert(doc_out, p_in):
    doc_out._body._body._insert_p(p_in._p)


def add_blank_lines(document, repeat=1, line_spacing=1.15):
    for _ in range(repeat):
        p = document.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.line_spacing = line_spacing
        p.paragraph_format.widow_control = True
        p.paragraph_format.space_after = Cm(0)

        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.right_indent = Cm(0)


def fix_hanging_prepositions(paragraph):
    for run in paragraph.runs:
        words = run.text.split()
        for i in range(1, len(words)):
            if words[i-1].lower() in ['а', 'в', 'и', 'на', 'не', 'по', 'с', 'г.', 'у', 'за', 'к', 'о', 'от', 'для', 'перед', 'при', 'через', 'со']:
                words[i - 1] = f"{words[i - 1]}\u00A0{words[i]}"  # добавляем неразрывный пробел перед текущим словом
                words[i] = ''
                i += 1
        run.text = ' '.join(filter(None, words))