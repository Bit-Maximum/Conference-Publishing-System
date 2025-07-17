import io
import os

from roman_arabic_numerals import conv

from docx.enum.section import WD_ORIENTATION
import copy
import re

from docx import Document
import docx
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from lxml import etree

from .markup import (add_blank_lines, set_align_center_paragraph, set_align_left_paragraph, fix_hanging_prepositions,
                     set_collection_markup, set_default_paragraph, set_title_paragraph, set_wide_title_paragraph)
from .styles import *


def compose_collection(path_to_file, data):
    output = compose_collection_imprint(path_to_file, data)

    for i, section in enumerate(data.get("Секции")):
        compose_collection_section_heading(output, i + 1, section, section.get("title"))
        if i < len(data) - 1:
            output.add_page_break()
    output.save(path_to_file)


def compose_collection_imprint(path_to_file, data):
    output = docx.Document()
    set_styles(output, path_to_file)
    output = docx.Document(path_to_file)
    set_collection_markup(output)
    compose_reverse_title_page(output, data)
    compose_contents_template(output)
    output.save(path_to_file)
    return output


def compose_reverse_title_page(doc, data):
    compose_indexes(doc, data.get("УДК"), data.get("ББК"))
    compose_editors(doc, data.get("Ред. коллегия"))
    compose_lib_link(doc, data.get("Библиография"), [item.get("ФИО") for item in data.get("Ред. коллегия")])
    compose_annotation(doc, data.get("Аннотация"))
    compose_requirements(doc)
    compose_publication_settings(doc, data.get("Дата публикации"))
    compose_publish_place(doc, data.get("Выпускные данные"))
    doc.add_page_break()


def compose_indexes(doc, udc, lbc):
    p = set_align_left_paragraph(doc, 1.0)
    run = p.add_run(f"УДК {udc}")
    index_style(run)
    p = set_align_left_paragraph(doc, 1.0)
    run = p.add_run(f"ББК {lbc}")
    subtext_picture_style(run)


def compose_editors(doc, editors):
    p = set_align_center_paragraph(doc, 1.0)
    run = p.add_run("Редакционная коллегия:")
    default_italic_style(run)
    for i in range(len(editors)):
        cur = f'{editors[i].get("ФИО")}, {editors[i].get("Должность")}' if editors[i].get("Должность") else editors[i].get(
            "ФИО")
        if i + 1 != len(editors):
            cur = cur + ';'
        p = set_align_center_paragraph(doc, 1.0)
        run = p.add_run(cur)
        index_style(run)
    add_blank_lines(doc, line_spacing=1.0)


def compose_lib_link(doc, lib_link, editors):
    title = lib_link.get("Заглавие")
    description = lib_link.get("Описание")
    city = lib_link.get("Город")
    date = lib_link.get("Дата")
    url = lib_link.get("URL")
    publishing = lib_link.get("Издательство")
    isbn = lib_link.get("ISBN") if lib_link.get("ISBN") != "" else "000-0-0000-0000-0"

    p = set_default_paragraph(doc, 1.0)
    run = p.add_run(f"{title}.\u00A0")
    collection_bold_style(run)

    run = p.add_run(
        f"{description}, {city}, {date} г. : материалы конференции / ред. кол.: {', '.join(editors)}. – {publishing.get('Город')} : {publishing.get('Название')}, {publishing.get('Год')}. - [")
    index_style(run)
    run = p.add_run("200 c.")
    mark_style(run)
    run = p.add_run("]. – ISBN\u00A0")
    index_style(run)
    run = p.add_run(isbn)
    mark_style(run)
    run = p.add_run(". – DOI\u00A0")
    index_style(run)
    run = p.add_run("DOI - ССЫЛКА")
    mark_style(run)
    run = p.add_run(". – URL:\u00A0")
    index_style(run)
    run = p.add_run(f"{url}")
    mark_style(run)
    run = p.add_run(". – Дата публикации:\u00A0")
    index_style(run)
    run = p.add_run(f"01.01.{publishing.get('Год')}")
    mark_style(run)
    run = p.add_run(". – Текст. Изображение : электронные.")
    index_style(run)
    add_blank_lines(doc, line_spacing=1.0)


def compose_annotation(doc, annotation):
    for paragraph in annotation:
        p = set_default_paragraph(doc, 1.0)
        run = p.add_run(paragraph)
        index_style(run)
    add_blank_lines(doc, line_spacing=1.0)


def compose_requirements(doc):
    compose_min_system(doc)
    add_blank_lines(doc, line_spacing=1.0)
    compose_config(doc)
    add_blank_lines(doc, line_spacing=1.0)


def compose_min_system(doc):
    p = set_align_center_paragraph(doc, 1.0)
    run = p.add_run("Текстовое электронное издание")
    default_italic_style(run)
    add_blank_lines(doc, line_spacing=1.0)

    p = set_align_center_paragraph(doc, 1.0)
    run = p.add_run("Минимальные системные требования:")
    index_style(run)

    p = set_align_center_paragraph(doc, 1.0)
    run = p.add_run("Веб-браузер Internet Explorer версии 6.0 или выше,\nOpera Версии 7.0 или выше, Google Chrome 3.0 или выше.")
    index_style(run)


def compose_config(doc):
    p = set_align_center_paragraph(doc, 1.0)
    run = p.add_run("Компьютер с доступом к сети Интернет.")
    index_style(run)

    p = set_align_center_paragraph(doc, 1.0)
    run = p.add_run("Минимальные требования к конфигурации и операционной системе компьютера\nопределяются требованиями перечисленных выше программных продуктов. ")
    index_style(run)


def compose_publication_settings(doc, date):
    p = set_align_center_paragraph(doc, 1.0)
    run = p.add_run("Размещено на сайте\u00A0")
    index_style(run)
    p = set_align_center_paragraph(doc, 1.0)
    run = p.add_run(f"{date} г.")
    mark_style(run)

    p = set_align_center_paragraph(doc, 1.0)
    run = p.add_run("Объем\u00A0")
    index_style(run)
    p = set_align_center_paragraph(doc, 1.0)
    run = p.add_run("4,50 Мб")
    mark_style(run)
    add_blank_lines(doc, line_spacing=1.0)


def compose_publish_place(doc, publication):
    place = publication.get("Название")
    post_code = publication.get("Индекс")
    address = publication.get("Адрес")
    email = publication.get("E-mail")
    phone = publication.get("Телефон")
    copyrite = publication.get("Copyright")
    year = publication.get("Год")

    p = set_align_center_paragraph(doc, 1.0)
    run = p.add_run(f"{place}\n{post_code}, {address}.")
    index_style(run)
    add_blank_lines(doc, line_spacing=1.0)

    p = set_align_center_paragraph(doc, 1.0)
    run = p.add_run(f"E-mail: {email}")
    index_style(run)
    p = set_align_center_paragraph(doc, 1.0)
    run = p.add_run(f"Тел.: {phone}")
    index_style(run)
    add_blank_lines(doc, line_spacing=1.0)

    p = set_align_center_paragraph(doc, 1.0)
    run = p.add_run(f"© {copyrite}, {year}")
    index_style(run)


def compose_contents_template(doc):
    p = set_title_paragraph(doc)
    run = p.add_run("СОДЕРЖАНИЕ")
    collection_headings_style(run)
    add_blank_lines(doc)

    p = set_align_left_paragraph(doc)
    run = p.add_run("[МЕСТО ДЛЯ СОДЕРЖНИЯ]\nЧтобы вставить интерактивное содержание, перейдите в раздел «Ссылки» и выберете опцию «Оглавление» на ваше усмотрение.\n\nВсе материалы сборника автоматически добавятся в оглавление, после нажатия на кнопку «Обновить оглавление».")
    collection_bold_style(run)
    doc.add_page_break()


def compose_collection_section_heading(doc, path_to_file, section_number, section_title):
    p = set_wide_title_paragraph(doc)
    p.style = doc.styles["Heading 1"]
    run = p.add_run(f"Секция {conv.arab_rom(section_number)}\n{section_title}")
    section_headings_style(run, 16)
    fix_hanging_prepositions(p)
    insert_underline(doc)
    doc.save(path_to_file)


def insert_underline(doc):
    p = set_wide_title_paragraph(doc)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(0.0), leader=WD_TAB_LEADER.LINES)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(17.0), leader=WD_TAB_LEADER.LINES)
    p.add_run('\t')
