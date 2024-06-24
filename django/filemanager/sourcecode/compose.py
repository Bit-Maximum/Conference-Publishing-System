import io
import os

import copy
import random
import re

from docx import Document
import docx
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import CT_Numbering
from docx.styles.style import _ParagraphStyle
from lxml import etree

from .markup import *
from .styles import *
from .shapes import paragraph_contains_formula, paragraph_contains_image, collect_images, add_border, \
    remove_images_directory, paragraph_contains_numbering

from cms.models import ControlsPage


def compose_document(doc_in, doc_out, data, session_num, img_dir, fix_hanging=False, doc_type="thesis", path_to_file_out=None):
    try:
        imgs = collect_images(doc_in, img_dir, session_num)
        if doc_type != "collection":
            for img in imgs:
                add_border(img, img, 2)
        doc_in.seek(0)

        document = Document(io.BytesIO(doc_in.read()))
        if doc_type != "collection":
            output = docx.Document()
            set_default_style(output, doc_out)
            set_markup(output)
            compose_heading(output, data, doc_type=doc_type)
            compose_main(output, document, imgs, title=data.get("title"), fix_hanging=fix_hanging,
                         have_sources=(len(data.get("sources")) > 0), doc_type=doc_type)
            compose_sources(output, data.get("sources"))
        else:
            output = doc_out
            index = compose_collection_heading(output, document, data)
            copy_document(output, document, images=imgs, starting_index=index, fix_hanging=fix_hanging)

        if path_to_file_out:
            output.save(path_to_file_out)
        else:
            output.save(doc_out)
    except Exception as e:
        print(e)
        raise e
    finally:
        remove_images_directory(img_dir, session_num)


def compose_heading(document, data: dict, doc_type="thesis"):
    compose_authors(document, names=data.get("authors").get("names"), cities=data.get("authors").get("cities"))
    add_blank_lines(document)

    if data.get("adviser", None) is not None:
        compose_adviser(
            document,
            name=data.get("adviser").get("name"),
            degree=data.get("adviser").get("degree"),
            academic_title=data.get("adviser").get("academic_title"),
            job_title=data.get("adviser").get("job_title"),
            job=data.get("adviser").get("job"),
        )
        add_blank_lines(document)

    compose_title(document, title=data.get("title"), grant=data.get("grant"), fix_hanging=True)
    add_blank_lines(document)

    if doc_type == "article":
        compose_abstract(document, abstract=data.get("abstract"))
        compose_keywords(document, keywords=data.get("keywords"))
        add_blank_lines(document)


def compose_collection_heading(doc_out, doc_in, data: dict):
    compose_title(doc_out, title=data.get("title"), doc_type="collection", grant=data.get("grant"))
    add_blank_lines(doc_out)
    compose_collection_authors(doc_out, names=data.get("authors").get("names"), cities=data.get("authors").get("cities"))
    current_element_index = find_and_compose_adviser(doc_out, doc_in)
    add_blank_lines(doc_out)
    index = find_and_compose_annotation_and_keywords(doc_out, doc_in, current_element_index)
    if index < 0:
        compose_abstract(doc_out, abstract=data.get("abstract"))
        compose_keywords(doc_out, keywords=data.get("keywords"))
    add_blank_lines(doc_out)
    return index


def compose_authors(document, names, cities):
    indexing = len(set(cities)) != 1

    if indexing:
        p = set_align_right_paragraph(document)
        for i in range(len(names)):
            run = p.add_run(names[i])
            style_authors_names(run)
            run = p.add_run(str(i + 1))
            style_authors_names(run)
            run.font.superscript = True
            if i < len(names) - 1:
                run = p.add_run(", ")
                style_authors_names(run)

        p = set_align_right_paragraph(document)
        for i in range(len(cities)):
            run = p.add_run(cities[i])
            style_heading_text(run)
            run = p.add_run(str(i + 1))
            style_heading_text(run)
            run.font.superscript = True
            if i < len(names) - 1:
                run = p.add_run(", ")
                style_heading_text(run)
    else:
        p = set_align_right_paragraph(document)
        run = p.add_run(", ".join(names))
        style_authors_names(run)

        p = set_align_right_paragraph(document)
        run = p.add_run(cities[0])
        style_heading_text(run)


def compose_collection_authors(document, names, cities):
    indexing = len(set(cities)) != 1

    if indexing:
        p = set_align_right_paragraph(document)
        for i in range(len(names)):
            run = p.add_run(names[i])
            style_collection_authors_names(run)
            run = p.add_run(str(i + 1))
            style_collection_authors_names(run)
            run.font.superscript = True
            if i < len(names) - 1:
                run = p.add_run(",\n")
                style_collection_authors_names(run)

        p = set_align_right_paragraph(document)
        for i in range(len(cities)):
            run = p.add_run(cities[i])
            style_heading_text(run)
            run = p.add_run(str(i + 1))
            style_heading_text(run)
            run.font.superscript = True
            run = p.add_run(", ")
            style_heading_text(run)
    else:
        p = set_align_right_paragraph(document)
        run = p.add_run(",\n".join(names))
        style_collection_authors_names(run)

        p = set_align_right_paragraph(document)
        run = p.add_run(cities[0])
        style_heading_text(run)


def compose_adviser(document, name: str = None, degree: str = None, academic_title: str = None, job_title: str = None, job: str = None, fix_hanging: bool = False):
    if ControlsPage.objects.first().overwrite_academic_title:
        academic_title = academic_title if academic_title else "ученое звание отсутствует"

    if job_title and job:
        job_string = f", {job_title} {job}"
    elif job:
        job_string = f" {job}"
    elif job_title:
        job_string = f", {job_title}"
    else:
        job_string = None

    adviser_title = ', '.join([degree, academic_title])

    tun_text = f"{adviser_title}{job_string}" if job_string else adviser_title

    p = set_align_right_paragraph(document)
    run = p.add_run("Научный руководитель:")
    style_heading_text(run)

    p = set_align_right_paragraph(document)
    run = p.add_run(tun_text)
    style_heading_text(run)

    p = set_align_right_paragraph(document)
    run = p.add_run(name)
    style_heading_text(run)

    if fix_hanging:
        fix_hanging_prepositions(p)


def find_and_compose_adviser(doc_out, doc_in):
    prefix, suffix, index = find_adviser(doc_in)
    if index < 0:
        return -1

    p = set_align_right_paragraph(doc_out)
    run = p.add_run(f"Научный руководитель: {prefix}")
    style_heading_text(run)

    p = set_align_right_paragraph(doc_out)
    run = p.add_run(suffix)
    style_heading_text(run)
    return index


def find_adviser(doc_in):
    index = find_element_startswith(doc_in, "Научный руководитель")
    if index < 0:
        return None, None, index

    prefix_p = doc_in.elements[index + 1]
    adviser_prefix = ', '.join(prefix_p.text.split(', ')[:-1])
    suffix_p = doc_in.elements[index + 2]
    adviser_suffix = f"{prefix_p.text.split(', ')[-1]} {suffix_p.text}" if suffix_p.text != "" else prefix_p.text.split(', ')[-1]
    return adviser_prefix, adviser_suffix, index + 3


def find_and_compose_annotation_and_keywords(doc_out, doc_in, starting_index):
    index_annotation = find_element_startswith(doc_in, "Аннотация.", starting_index=starting_index)
    if index_annotation < 0:
        return -1

    annotation_p = doc_in.elements[index_annotation]
    p_out = set_default_paragraph(doc_out)
    for run_in in annotation_p.runs:
        run_out = p_out.add_run(run_in.text)
        copy_run_style(run_out, run_in)
        change_font_size(run_out, 14)

    index_keywords = find_element_startswith(doc_in, "Ключевые слова:", starting_index=index_annotation + 1)
    if index_keywords < 0:
        return -1

    keywords_p = doc_in.elements[index_keywords]
    p_out = set_default_paragraph(doc_out)
    for run_in in keywords_p.runs:
        run_out = p_out.add_run(run_in.text)
        copy_run_style(run_out, run_in)
        change_font_size(run_out, 14)
    return index_keywords + 1


def find_element_startswith(doc_in, text, starting_index=0, max_depth=15):
    size = len(doc_in.elements)
    i = starting_index
    while i < size and i < max_depth:
        try:
            element = doc_in.elements[i]
        except IndexError:
            return
        if isinstance(element, docx.text.paragraph.Paragraph):
            if element.text.startswith(text) and i + 2 < size:
                return i
        i += 1
    return -1


def skip_empty_paragraphs(doc_in, starting_index=0, max_depth=20):
    size = len(doc_in.elements)
    i = starting_index
    while i < size:
        try:
            element = doc_in.elements[i]
        except IndexError:
            return

        if isinstance(element, docx.text.paragraph.Paragraph):
            if not is_paragraph_empty(element):
                return i
        i += 1
    return starting_index


def old_compose_adviser(document, data, fix_hanging: bool = False):
    info = data.get("adviser")
    degree = info.get("degree")
    academic_title = info.get("academic_title") if info.get("academic_title") else "ученое звание отсутствует"
    job_title = info.get("job_title")
    job = info.get("job")
    name = info.get("name")

    job_string = f"{job_title} {job} {name}" if job_title and job else f"{job} {name}"

    p = set_align_right_paragraph(document)
    run = p.add_run(f"Научный руководитель: {degree}, {academic_title}, {job_string}")
    style_heading_text(run)
    if fix_hanging:
        fix_hanging_prepositions(p)


def compose_title(document, title, doc_type="thesis", grant=None, fix_hanging: bool = True):
    p = set_title_paragraph(document)
    run = p.add_run(title)
    style_title(run, font_size=15)
    if doc_type == "collection":
        p.style = document.styles['Heading 2']

    if fix_hanging:
        fix_hanging_prepositions(p)
    if grant:
        insert_footnote(p, grant)


def compose_abstract(document, abstract):
    p = set_default_paragraph(document)
    run = p.add_run("Аннотация.\u00A0")
    style_abstract_or_keywords(run)

    run = p.add_run(f"{abstract}")
    reset_run_style(run)
    run.italic = False


def compose_keywords(document, keywords):
    p = set_default_paragraph(document)
    run = p.add_run("Ключевые слова:\u00A0")
    style_abstract_or_keywords(run)

    run = p.add_run(f"{keywords}")
    reset_run_style(run)
    run.italic = False


def compose_sources(document, sources=None, fix_hanging: bool = False):
    if not sources:
        return

    add_blank_lines(document)
    p = set_align_center_paragraph(document)
    run = p.add_run("Список литературы")
    reset_run_style(run)
    run.font.bold = True

    for source in sources:
        source = source.strip()  # TODO remove when redeploy because source already filters when they added
        if bool(re.match(r"^\d+[.)] ", source)):  # TODO remove when redeploy
            source = ' '.join(source.split()[1:])  # TODO remove when redeploy

        source = source.replace("\n", "")  # TODO remove when redeploy
        source = source.replace("\r", "")  # TODO remove when redeploy

        p = set_source_paragraph(document)
        p.style = 'List Number'
        p.paragraph_format.tab_stops.add_tab_stop(Cm(1.75), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(2), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.left_indent = Cm(0)
        run = p.add_run(source)
        reset_run_style(run)
        if fix_hanging:
            fix_hanging_prepositions(p)


def compose_picture_text(text: str, img_num: int):
    if text is None or not text.strip():
        return f"Рис. {img_num + 1}", "."

    if text.lower().startswith(("рис. ", "рисунок ")) or text.lower().startswith(("рис.")):
        words = text.split(' ')
        if words[0].lower() == "рис.":
            words[0] = "Рис."
            words[1] = words[1].rstrip('.')
            if words[1] != str(img_num + 1):
                words[1] = f"{img_num + 1}"
            prefix = " ".join(words[:2])
            subtext = f'. {" ".join(words[2:])}'
        else:
            tags = text.split('.')
            if tags[1].isdigit():
                words[0] = f"Рис. {img_num + 1}"
            prefix = " ".join(words[:1])
            subtext = f'. {" ".join(words[1:])}'
    else:
        prefix = f"Рис. {img_num + 1}"
        subtext = f". {text}"
    subtext = subtext.rstrip().rstrip(".")
    return prefix, subtext


def compose_picture_text_old(text: str, img_num: int):
    if text.lower().startswith(("рис. ", "рисунок ")):
        words = text.split()
        words[0] = "Рис."
        words[1] = words[1].rstrip('.')
        if words[1] != str(img_num + 1):
            words[1] = f"{img_num + 1}"
        prefix = " ".join(words[:2])
        subtext = f'. {" ".join(words[2:])}'
    else:
        prefix = f"Рис. {img_num + 1}"
        subtext = f". {text}"
    subtext = subtext.rstrip().rstrip(".")
    return prefix, subtext


def compose_single_picture_text(text: str):
    if text.lower().startswith(("рис. ", "рисунок ")):
        words = text.split()
        subtext = " ".join(words[2:])
    else:
        subtext = text
    subtext = subtext.rstrip().rstrip(".")
    return subtext


def compose_table_text(text, table_num):
    if text.lower().startswith(("таб. ", "таблица ")):
        words = text.split()
        words[0] = "Таблица"
        words[1] = words[1].rstrip('.')
        if words[1] != str(table_num + 1):
            words[1] = str(table_num + 1)
        prefix = " ".join(words[:2])
        subtext = " ".join(words[2:])
    else:
        prefix = f"Таблица {table_num + 1}"
        subtext = text
    subtext = subtext.rstrip().rstrip(".")
    return prefix, subtext


def is_table_prefix_on_top(document, index):
    steps_left = 2
    j = index
    while steps_left:
        j -= 1
        temp = document.elements[j]
        if isinstance(temp, docx.text.paragraph.Paragraph):
            if is_paragraph_empty(temp):
                continue

            if temp.text.strip().lower().startswith(("таб.", "таблица")):
                if steps_left == 1:
                    steps_left -= 1
                    return True, j
                return None, j
            steps_left -= 1
        else:
            return False, j
    return False, j


def check_if_sources_next(paragraph):
    if not isinstance(paragraph, docx.text.paragraph.Paragraph):
        return False
    return paragraph.text.strip().lower() == "список литературы"
    # return paragraph.text.strip().lower() in ["список литературы", "список источников"]


def delete_paragraph(paragraph):
    paragraph.clear()
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def compose_main(doc_out, doc_in, imgs, title=None, fix_hanging: bool = False, have_sources: bool = False, doc_type="thesis"):
    tables = doc_in.tables
    table_counter = 0
    image_counter = 0
    list_number_counter = 0
    last_list_number_pointer = None
    size = len(doc_in.elements)
    i = 0
    while i < size:
        try:
            element = doc_in.elements[i]
        except IndexError:
            return

        if isinstance(element, docx.table.Table):
            prefix_on_top, index_on_top = is_table_prefix_on_top(doc_in, i)
            if prefix_on_top:
                index = index_on_top + 1
                delete_paragraph(doc_out.elements[-3])
                delete_paragraph(doc_out.elements[-2])
            else:
                if prefix_on_top is None:
                    delete_paragraph(doc_out.elements[-2])  # Если перед таблицей есть префикс, но подпись под таблицей
                i += 1
                index = i

            add_blank_lines(doc_out)

            template = tables[table_counter]
            new_tbl = get_table_copy(template)

            headings = doc_in.elements[index].text
            prefix, t_name = compose_table_text(headings, table_counter)

            p_out = set_align_right_paragraph(doc_out)
            run = p_out.add_run(prefix)
            set_style_table_prefix(run)
            if doc_type == "article" or doc_type == "collection":
                run.font.bold = True

            p_out = set_align_center_paragraph(doc_out)
            run = p_out.add_run(t_name)
            set_style_table_name(run)

            fix_hanging_prepositions(p_out)

            p_out._p.addnext(new_tbl)

            t_out = doc_out.tables[table_counter]
            set_style_default_table(t_out)

            add_blank_lines(doc_out)

            table_counter += 1
            i += 1
            continue

        if isinstance(element, docx.section.Section):
            i += 1
            continue

        if isinstance(element, docx.text.paragraph.Paragraph):
            if paragraph_contains_formula(element):
                add_blank_lines(doc_out)
                direct_insert(doc_out, element)
                element = doc_in.elements[i]
                if element.text != "":
                    p_out = set_align_center_paragraph(doc_out)
                    for run_in in element.runs:
                        run_out = p_out.add_run(re.sub(r" +", " ", run_in.text))
                        run_default_style(run_out, run_in)
                add_blank_lines(doc_out)
                i += 1
                continue

            if paragraph_contains_image(element):
                add_blank_lines(doc_out)
                p_out = set_align_center_paragraph(doc_out)
                run = p_out.add_run()
                img_num = len(imgs) - 1 if image_counter >= len(imgs) else image_counter
                run.add_picture(imgs[img_num], width=Cm(10))
                i += 1
                element = doc_in.elements[i]
                if not isinstance(element, docx.text.paragraph.Paragraph):
                    i += 1
                    continue

                if is_paragraph_empty(element):
                    element = doc_in.elements[i - 1]  # Если под изображением нет подписи, то скорее всего он на параграфе с изображением

                p_out = set_align_center_paragraph(doc_out)
                prefix, subtext = compose_picture_text(element.text, image_counter)

                run_out = p_out.add_run(prefix)
                prefix_picture_style(run_out)

                run_out = p_out.add_run(subtext)
                subtext_picture_style(run_out)

                fix_hanging_prepositions(p_out)
                add_blank_lines(doc_out)
                image_counter += 1
                i += 1
                continue

            if is_paragraph_empty(element) and i + 1 < size:
                i += 1
                if have_sources and check_if_sources_next(doc_in.elements[i]):
                    break
                continue

            if i == 0:
                if title.lower() in element.text.lower():
                    i += 1
                    continue

            p_out = doc_out.add_paragraph()

            if paragraph_contains_numbering(element):
                list_number_counter += 1
                last_list_number_pointer = p_out
                p_out.style = 'List Bullet'
                p_out.paragraph_format.tab_stops.add_tab_stop(Cm(1.75), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)

                paragraph_default_style(p_out, element)
                # p_out.paragraph_format.first_line_indent = Cm(1.25)
                # p_out.paragraph_format.left_indent = Cm(0)
            else:
                paragraph_default_style(p_out, element)

            for run_in in element.runs:
                run_out = p_out.add_run(re.sub(r" +", " ", run_in.text))
                run_default_style(run_out, run_in)

            if fix_hanging:
                fix_hanging_prepositions(p_out)
            i += 1
            continue
        else:
            i += 1

    if list_number_counter == 1:
        # Если только один элемент в списке, то предполагаем, что списком оформили
        # только первый параграф и далее есть остальные, начинающиеся с 2)

        last_list_number_pointer.style = None
        last_list_number_pointer.runs[0].text = f"1) {last_list_number_pointer.runs[0].text}"


def copy_document(doc_out, doc_in, images, starting_index=0, fix_hanging: bool = False):
    tables = doc_in.tables
    table_counter = 0
    image_counter = 0
    list_number_counter = 0
    sources_found = False
    size = len(doc_in.elements)
    i = skip_empty_paragraphs(doc_in, starting_index=starting_index)

    while i < size:
        try:
            element = doc_in.elements[i]
        except IndexError:
            return

        if isinstance(element, docx.table.Table):
            previous = doc_out.paragraphs[-1]
            for run in previous.runs:
                run.font.bold = True

            template = tables[table_counter]
            new_tbl = get_table_copy(template)
            previous._p.addnext(new_tbl)

            t_out = doc_out.tables[-1]
            set_style_default_table(t_out, font_size=12)

            table_counter += 1
            i += 1
            continue

        if isinstance(element, docx.section.Section):
            i += 1
            continue

        if isinstance(element, docx.text.paragraph.Paragraph):
            if paragraph_contains_formula(element):
                direct_insert(doc_out, element)
                element = doc_in.elements[i]
                if element.text != "":
                    p_out = set_align_center_paragraph(doc_out)
                    for run_in in element.runs:
                        run_out = p_out.add_run(re.sub(r" +", " ", run_in.text))
                        run_default_style(run_out, run_in)
                i += 1
                continue

            if paragraph_contains_image(element):
                p_out = set_align_center_paragraph(doc_out)
                run = p_out.add_run()
                img_num = len(images) - 1 if image_counter >= len(images) else image_counter
                run.add_picture(images[img_num], width=Cm(10))
                i += 1
                element = doc_in.elements[i]
                if not isinstance(element, docx.text.paragraph.Paragraph):
                    i += 1
                    continue

                if is_paragraph_empty(element):
                    element = doc_in.elements[i - 1]  # Если под изображением нет подписи, то скорее всего он на параграфе с изображением

                p_out = doc_out.add_paragraph()
                copy_paragraph_style(p_out, element)

                for run_in in element.runs:
                    run_out = p_out.add_run(re.sub(r" +", " ", run_in.text))
                    copy_run_style(run_out, run_in)

                image_counter += 1
                i += 1
                continue

            print(str(element.style))
            if paragraph_contains_numbering(element):
                if sources_found or str(element.style).split("'")[1] == 'List Number':
                    list_number_counter += 1
                    p_out = doc_out.add_paragraph()
                    copy_paragraph_style(p_out, element)
                    run = p_out.add_run(f"{list_number_counter}.\u00A0")
                    run_default_style(run)
                    change_font_size(run, 14)
                    for run_in in element.runs:
                        run_out = p_out.add_run(re.sub(r" +", " ", run_in.text))
                        copy_run_style(run_out, run_in)
                        change_font_size(run_out, 14)
                else:
                    p_out = doc_out.add_paragraph()
                    copy_paragraph_style(p_out, element)
                    p_out.style = 'List Bullet'
                    p_out.paragraph_format.tab_stops.add_tab_stop(Cm(1.75), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
                    for run_in in element.runs:
                        run_out = p_out.add_run(re.sub(r" +", " ", run_in.text))
                        copy_run_style(run_out, run_in)
                        change_font_size(run_out, 14)

                    # TODO: Закоментирован старый подход. Direct_insert вставлял параграф,
                    #       однако случайным образом мог пропустить следующий за ним.
                    #       В связи с этим было решено использовать подход как и при первоначальной вёрстке документов.

                    # direct_insert(doc_out, element)
                    # print(doc_out.paragraphs[-1].text)

                    # p_out = doc_out.paragraphs[-1]
                    # for run_out in element.runs:
                    #     change_font_size(run_out, 14)
                    # print(doc_in.elements[i + 1].text)
                i += 1
                continue

            p_out = doc_out.add_paragraph()
            copy_paragraph_style(p_out, element)
            if check_if_sources_next(element):
                p_out.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                p_out.paragraph_format.first_line_indent = Cm(1.25)
                sources_found = True

            for run_in in element.runs:
                run_out = p_out.add_run(re.sub(r" +", " ", run_in.text))
                copy_run_style(run_out, run_in)
                change_font_size(run_out, 14)

            i += 1
            continue
        i += 1
